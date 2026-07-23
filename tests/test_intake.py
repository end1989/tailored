"""Tests for backend/app/services/intake.py."""
from __future__ import annotations

import io

import docx
import pypdf

from backend.app.schemas import Contact, MasterProfile, UsageInfo
from backend.app.services.intake import (
    INTAKE_SYSTEM,
    IntakeResult,
    build_master_profile,
    extract_text,
)


def test_extract_text_txt_bytes():
    kind, text = extract_text(
        "notes.txt", "Jordan Rivera\nSenior engineer notes.".encode("utf-8")
    )
    assert kind == "txt"
    assert text == "Jordan Rivera\nSenior engineer notes."


def test_extract_text_unknown_extension_defaults_to_txt():
    kind, text = extract_text("bio.md", b"# Bio\nBackend engineer.")
    assert kind == "txt"
    assert text == "# Bio\nBackend engineer."


def test_extract_text_docx_roundtrip():
    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Jordan Rivera")
    document.add_paragraph("Senior Software Engineer at Cascade Analytics")
    document.save(buffer)
    kind, text = extract_text("resume.docx", buffer.getvalue())
    assert kind == "docx"
    assert "Jordan Rivera" in text
    assert "Senior Software Engineer at Cascade Analytics" in text


def test_extract_text_pdf_via_stubbed_reader(monkeypatch):
    class StubPage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class StubReader:
        def __init__(self, stream) -> None:
            self.pages = [StubPage("Page one text"), StubPage("Page two text")]

    monkeypatch.setattr(pypdf, "PdfReader", StubReader)
    kind, text = extract_text("resume.pdf", b"%PDF-1.7 fake bytes")
    assert kind == "pdf"
    assert text == "Page one text\nPage two text"


def test_intake_result_shape():
    assert set(IntakeResult.model_fields) == {"contact", "master_profile"}


def test_build_master_profile_returns_fixture_backed_profile(claude_fake):
    docs = [
        "JORDAN RIVERA\nSenior Software Engineer\n...full resume text...",
        "Extra career notes",
    ]
    profile, contact, usage = build_master_profile(docs, claude_fake)
    assert isinstance(profile, MasterProfile)
    assert isinstance(contact, Contact)
    assert contact.name == "Jordan Rivera"
    assert [e.company for e in profile.experiences] == [
        "Cascade Analytics",
        "Brightline Software",
    ]
    assert all(len(e.bullets) == 4 for e in profile.experiences)
    assert all(b.tags for e in profile.experiences for b in e.bullets)
    assert usage == UsageInfo(input_tokens=0, output_tokens=0, cost_usd=0.0)


def test_build_master_profile_records_intake_call(claude_fake):
    build_master_profile(["Doc A body", "Doc B body"], claude_fake)
    call = claude_fake.calls[-1]
    assert call["task"] == "intake"
    assert call["schema_model_name"] == "IntakeResult"
    assert call["tools"] is None
    assert "--- DOCUMENT 1 ---" in call["user_content"]
    assert "Doc A body" in call["user_content"]
    assert "--- DOCUMENT 2 ---" in call["user_content"]
    assert "Doc B body" in call["user_content"]
    assert call["system"] == INTAKE_SYSTEM
    assert "NEVER invent" in INTAKE_SYSTEM
    assert "VERBATIM" in INTAKE_SYSTEM
